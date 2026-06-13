from docx import Document

# Load the document
doc = Document(r"C:\Users\admin\Downloads\SA SIL - Kylie Furgusson.docx")

print("=" * 70)
print("📄 SERVICE AGREEMENT TEMPLATE STRUCTURE")
print("=" * 70)
print()

# Extract all paragraphs with their styles
print("DOCUMENT SECTIONS & CONTENT:")
print("-" * 70)

section_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if not text:
        continue
    
    # Identify headings
    is_heading = para.style.name.startswith('Heading')
    
    if is_heading:
        section_count += 1
        print(f"\n[Section {section_count}] {text}")
        print("  " + "-" * 60)
    else:
        # Show first 100 chars of paragraph
        preview = text[:100] + "..." if len(text) > 100 else text
        print(f"  • {preview}")

print()
print()
print("=" * 70)
print(f"📊 DOCUMENT STATISTICS")
print("=" * 70)
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")
print(f"Total Sections: {section_count}")

print()
print("TABLES:")
print("-" * 70)
for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    print(f"  Rows: {len(table.rows)}")
    print(f"  Columns: {len(table.columns)}")
    if len(table.rows) > 0:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"  Headers: {headers}")
